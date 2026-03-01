#!/usr/bin/env python3
"""
Generate comprehensive O-1 visa CV as a Word document for Anekha Sokhal.
Merges content from: 6-page Fulbright CV, May 2025 resume, and portfolio website.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Constants ────────────────────────────────────────────────────────────────

DRIVE_PLACEHOLDER = 'https://drive.google.com/PLACEHOLDER'

# ── Helpers ──────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, text, url, color='0563C1'):
    """Add a clickable hyperlink to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt in half-points
    rPr.append(sz)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_evidence_links(paragraph, website_url=None):
    """Append ' [Drive] [Website]' as hyperlinks at the end of a paragraph.

    All claims get a [Drive] link (placeholder). Only claims with a public URL
    also get a [Website] link.
    """
    run = paragraph.add_run('  ')
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_hyperlink(paragraph, '[Drive]', DRIVE_PLACEHOLDER, color='E67E22')
    if website_url:
        spacer = paragraph.add_run(' ')
        spacer.font.size = Pt(10.5)
        spacer.font.name = 'Calibri'
        add_hyperlink(paragraph, '[Website]', website_url, color='27AE60')


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a thin horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_section_heading(doc, text):
    """Add a section heading: bold, 13pt, uppercase, with underline-like bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F2937')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_role_header(doc, company, title, location, dates):
    """Add a role header: Bold company | Italic title | right-aligned dates."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    # Company name (bold)
    run_company = p.add_run(company)
    run_company.bold = True
    run_company.font.size = Pt(11)
    run_company.font.name = 'Calibri'

    if location:
        run_loc = p.add_run(f'  —  {location}')
        run_loc.font.size = Pt(11)
        run_loc.font.name = 'Calibri'

    p.add_run('\n')

    # Title (italic)
    run_title = p.add_run(title)
    run_title.italic = True
    run_title.font.size = Pt(11)
    run_title.font.name = 'Calibri'

    # Dates (right side - use tab stop)
    run_dates = p.add_run(f'\t{dates}')
    run_dates.font.size = Pt(11)
    run_dates.font.name = 'Calibri'

    # Add a right-aligned tab stop
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    return p


def add_sub_header(doc, text):
    """Add a sub-header within a role (e.g., 'Line of business:', 'Main Duties:')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point at the specified indent level."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    # Clear the default run and add our formatted one
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 + level * 0.3)
    return p


def add_sub_bullet(doc, text):
    """Add a sub-bullet (indented, with dash marker)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.75)
    run = p.add_run(f'– {text}')
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p


def add_body_text(doc, text):
    """Add a regular paragraph of body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p


def add_education_entry(doc, institution, degree, dates, honors=None, details=None):
    """Add an education entry."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

    run_inst = p.add_run(institution)
    run_inst.bold = True
    run_inst.font.size = Pt(11)
    run_inst.font.name = 'Calibri'

    run_dates = p.add_run(f'\t{dates}')
    run_dates.font.size = Pt(11)
    run_dates.font.name = 'Calibri'

    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    if degree:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        run_deg = p2.add_run(degree)
        run_deg.italic = True
        run_deg.font.size = Pt(10.5)
        run_deg.font.name = 'Calibri'

    if honors:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        run_h = p3.add_run(f'Honors: {honors}')
        run_h.font.size = Pt(10.5)
        run_h.font.name = 'Calibri'
        run_h.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    if details:
        for detail in details:
            add_bullet(doc, detail)


# ── Main Document ────────────────────────────────────────────────────────────

def generate_cv():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # ── Configure List Bullet style ──
    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Calibri'
        lb.font.size = Pt(10.5)

    # ════════════════════════════════════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════════════════════════════════════
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run('ANEKHA SOKHAL')
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = 'Calibri'
    name_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(2)
    contact_run = contact_p.add_run(
        'anekhas@gmail.com  |  +1 832 677 5084  |  linkedin.com/in/anekha  |  anekha.github.io'
    )
    contact_run.font.size = Pt(10.5)
    contact_run.font.name = 'Calibri'
    contact_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    location_p = doc.add_paragraph()
    location_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location_p.paragraph_format.space_before = Pt(0)
    location_p.paragraph_format.space_after = Pt(4)
    loc_run = location_p.add_run('Houston, Texas')
    loc_run.font.size = Pt(10.5)
    loc_run.font.name = 'Calibri'
    loc_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    add_horizontal_line(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 1. PROFESSIONAL SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Professional Summary')
    add_body_text(
        doc,
        'Fulbright Scholar, AI/ML Engineer & Founder building AI for financial markets (Moshi). '
        'Published researcher (ICLR 2026) with NASA, Stability AI, and Rice University. '
        'Former JP Morgan derivatives trader with 7 years in global energy markets. '
        'Track record spanning AI/ML, quantitative finance, entrepreneurship, and social impact '
        'across the US, UK, and emerging markets.'
    )

    # ════════════════════════════════════════════════════════════════════════
    # 2. EDUCATION
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Education')

    # Rice University
    add_education_entry(
        doc,
        'Rice University, Houston, TX',
        'Professional Master of Data Science — Machine Learning & Computer Vision Specialization',
        '2023 – 2025',
        honors='Fulbright Elsevier Data Analytics Award (2023) | IBM AI Fellowship Nominee (2024) | Women in AI North America Finalist (2025)',
    )
    # Evidence links for Rice education
    p_rice_edu = add_bullet(doc, 'Fulbright profile at Rice')
    add_evidence_links(p_rice_edu, website_url='https://fulbright.rice.edu/people/anekha-sokhal')
    p_rice_eng = add_bullet(doc, 'Rice Engineering feature')
    add_evidence_links(p_rice_eng, website_url='https://engineering.rice.edu/news/fulbrighters-engineering-professional-masters-program-rice-home-away-home')

    # University of Warwick
    add_education_entry(
        doc,
        'The University of Warwick, UK',
        'BSc (Hons) Mathematics, Operational Research, Statistics, Economics (MORSE) — 2:1 obtained',
        'Oct 2011 – Jun 2014',
        details=[
            '3rd Year: Bayesian Statistics and Decision Theory, Multivariate Statistics, Bayesian Forecasting and Intervention, '
            'Principles of Finance I, Operational Research for Strategic Planning, Decision Analysis, The Practice of Operational Research, '
            'Topics in Economic Theory, Combinatorics',
            '2nd Year: Stochastic Processes, Mathematical Methods, Mathematical Statistics Part A, Mathematical Statistics Part B, '
            'Linear Statistical Modelling, Foundations of Finance, Foundations of Accounting, Mathematical Programming II, Simulation, '
            'Mathematical Economics 1A, Mathematical Economics 1B',
            '1st Year: Linear Algebra, Analysis, Introduction to Quantitative Economics, Mathematical Programming I, '
            'Statistical Laboratory I, Applications of Algebra and Analysis, Statistical Computing, Introduction to Probability, '
            'Differential Equations, Foundations',
        ],
    )

    # Queen Elizabeth College
    add_education_entry(
        doc,
        'Queen Elizabeth College, Cambridge International Examinations',
        None,
        'Jan 2009 – Nov 2010',
        details=[
            'A-Level Subjects: Mathematics A*, Chemistry A*, Physics A*; AS Level Subjects: French A, General Paper A',
        ],
    )

    # Sodnac State Secondary School
    add_education_entry(
        doc,
        'Sodnac State Secondary School, Cambridge International Examinations',
        None,
        'Jan 2007 – Nov 2008',
        details=[
            'O-Level Subjects: 8 A(a) grades — Mathematics, Additional Mathematics, Physics, Chemistry, Economics, Accounts, English, French',
        ],
    )

    # ════════════════════════════════════════════════════════════════════════
    # 3. PROFESSIONAL QUALIFICATIONS & CERTIFICATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Professional Qualifications & Certifications')
    add_bullet(doc, 'Chartered Institute for Securities & Investment (CISI): Securities, Derivatives and UK Regulations Papers (Oct 2014)')
    add_bullet(doc, 'Certifications: AWS, Modal, SQL, Spark, Hadoop')
    add_bullet(doc, 'Le Wagon Data Science Bootcamp (Dec 2022 – Jun 2023)')

    # ════════════════════════════════════════════════════════════════════════
    # 4. PROFESSIONAL EXPERIENCE
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Professional Experience')

    # ── Moshi ──
    add_role_header(doc, 'Moshi', 'Founder & Lead ML Engineer', 'Houston, Texas', 'Jun 2025 – Present')
    p = add_bullet(doc, 'VC-backed startup building 24/7 agentic quant analysts for financial markets')
    add_evidence_links(p, website_url='https://www.moshi.team/')
    add_bullet(doc, 'Live with institutional, fund, and corporate pilot customers')
    add_bullet(doc, 'Built and led a team of 5 in 6 months')

    # ── JewelVision ──
    add_role_header(doc, 'JewelVision', 'Founder & Lead AI Engineer', 'Houston, Texas', 'Jan 2024 – Dec 2024')
    p = add_bullet(doc, 'Selected as 1 of 8 startups for Rice\'s Liu Idea Lab Summer Venture Studio, a competitive accelerator for student founders')
    add_evidence_links(p, website_url='https://jewelvision.ai')
    add_bullet(doc, 'Spearheaded development of 3D asset generation and AR virtual try-on tools using multimodal AI (vision, LLMs, AR/VR)')
    add_bullet(doc, 'Raised $18K in funding and drove a cross-functional technical team; delivered live demos to 50+ clients and investors')
    p = add_bullet(doc, 'Pitched at JewelVision Demo Day to 100+ investors and researchers')
    add_evidence_links(p, website_url='https://www.youtube.com/live/WCl0LXecB_Y?si=emex9cpeQKH6F6dd')
    # Lilie feature links
    p = add_bullet(doc, 'Lilie feature — Rice News')
    add_evidence_links(p, website_url='https://news.rice.edu/news/2024/rice-entrepreneurship-lab-unveils-teams-2024-lilie-summer-venture-studio-accelerator')
    p = add_bullet(doc, 'Lilie feature — InnovationMap')
    add_evidence_links(p, website_url='https://houston.innovationmap.com/rice-university-liu-idea-lab-2024-2668479547.html')
    p = add_bullet(doc, 'Lilie feature — Rice News (Startups)')
    add_evidence_links(p, website_url='https://news.rice.edu/news/2024/lilie-showcases-8-student-startups-are-making-difference-their-industries')
    p = add_bullet(doc, 'Lilie feature — Rice Entrepreneurship Blog')
    add_evidence_links(p, website_url='https://entrepreneurship.rice.edu/blog/meet-the-2024-summer-venture-studio-cohort')

    # ── ESKA International ──
    add_role_header(doc, 'ESKA International', 'Head of Strategy & Data Analytics / Chief Operating Officer (COO)', 'UK', 'Apr 2020 – Apr 2023')
    p = add_body_text(
        doc,
        'Line of business: After one year of working part-time, moved to a full-time internal role. '
        'ESKA is a vertically integrated business with retail and wholesale operations in the UK and a unique '
        'manufacturing partner company based in Mauritius that cuts and polishes gemstones self-sourced from across '
        'the world locally and manufactures high-end jewellery to exceptional quality for consumer markets under the '
        'supervision of a Gemmologist.'
    )
    add_evidence_links(p, website_url='https://legemjewels.com')
    add_sub_header(doc, 'Main Duties:')
    add_bullet(doc, 'Collaborated with the CEO to support diverse business operations across all core areas of the business, '
               'including Product & Operations (gemstone trading, cutting, jewellery manufacturing), Finance, Marketing, Strategy '
               'and Consumer & Sales and establishing strong cohesion amongst cross-functional teams within these areas')
    add_bullet(doc, 'Monitored KPIs for production, customers, and sales. Partnered with management to budget for sufficient '
               'investment capital to achieve growth targets over the near term')
    add_bullet(doc, 'Researched investment opportunities, produced short and long-term cash flow forecast models, revenue analysis, '
               'business opportunities and presented proposals to the board of shareholders. Managed capital investment and expenses '
               'to ensure the company achieved targets relative to growth and profitability')
    add_bullet(doc, 'Oversaw operations and employee productivity, building a highly inclusive culture, ensuring team members '
               'thrived and organisational outcomes were met')
    add_bullet(doc, 'Adhered to company, state, and local business requirements, enforcing compliance and taking action when necessary')

    add_sub_header(doc, 'Key Achievements:')
    add_bullet(doc, 'Product and Operations: Analysed internal operations and identified areas of process enhancement. Drove the '
               'change to a more technology-based business and transition to a self-sufficient manufacturing house by investing in '
               'CAD design, casting machines and printing machines')
    add_bullet(doc, 'Finance: Initiated and led the migration of all financial accounts to professional accounting software, '
               'including full account reconciliation, completed with sign off from professional accountants and bookkeepers')
    add_bullet(doc, 'Marketing: Responsible for the entire branding campaign. Development and action of integrated marketing '
               'communication plans and budgets; trained and managed team members; understanding both B2B and B2C marketing; '
               'led the creative design and internal and external communications, including website, social media, PR, trade, and '
               'community engagement. Targeted product marketing after collecting data from customer surveys and incorporating '
               'feedback into marketing campaigns')
    add_bullet(doc, 'Strategy: Designed and translated strategy into actionable goals for performance and growth, helping to implement '
               'organisation-wide goal setting and operational planning through market research, internal interviews, and business analytics')
    add_bullet(doc, 'Customer and Sales: Transitioned to a customer-centric business using historical customer data to identify trends '
               'and provide a more personalised service to existing customers. Widened the customer base by utilising multiple selling '
               'channels, increasing brand awareness, and meeting customer need through hands-on account management. Collaborating '
               'with all team members across sales, product, and operations to meet customer deadlines on time')
    add_bullet(doc, 'Drove strategy across Europe & Africa, cutting costs by 17% via cloud migration, workflow reform, and staff training')
    add_bullet(doc, 'Launched an e-commerce site and doubled online engagement QoQ via rebranding, A/B testing, and analytics-led campaigns')
    add_bullet(doc, 'Led international sales expansion, closing the year\'s largest deal and generating 20% of annual revenue')

    # ── JP Morgan — Derivatives Trader ──
    add_role_header(doc, 'JP Morgan', 'Derivatives Trader, Energy Derivatives Markets', 'London, UK', 'Jan 2017 – Jun 2021')
    add_body_text(
        doc,
        'Line of business: Trader at JP Morgan in Commodities, based in London and working with Sydney, Singapore, '
        'New York, and Houston.\n'
        'Reporting: Team of two Traders, reporting to the Head of Coal, Gas and Emissions within Global Energy Trading.'
    )
    add_sub_header(doc, 'Main Duties:')
    add_bullet(doc, 'Primary Trader in London on the gas derivatives book (UK, EU, Asia, and US natural gas), managing a $60m '
               'portfolio of vanilla and exotic risk using financial instruments. The book had exposure to pipeline gas, LNG, '
               'carbon emissions, FX, and Rates risk')
    add_bullet(doc, 'Managing gas risk during the green energy transition as the energy sector transformed from fossil-based to '
               'low-carbon. The past decade has seen a significant rise in the importance of gas as a fuel which has generated a '
               'vast inflow of interest in gas, carbon, and renewable energy markets from a range of investors where I presented '
               'myself as a key market person')
    add_bullet(doc, 'Developed a data-based system to successfully identify market opportunities and trade recommendations using '
               'a combination of fundamental research, proprietary analytical methodologies, and quantitative analysis to drive views')
    add_bullet(doc, 'Competitive market-making to a broad client base in the blue-chip investor space, providing liquidity under '
               'all market conditions')
    add_bullet(doc, 'Partnered with Marketers & Research analysts to provide thorough coverage to a client. Developed structured '
               'deals for client hedge programmes through discussing pricing alternatives using a combination of derivatives and '
               'establishing a hedging strategy')

    add_sub_header(doc, 'Key Achievements:')
    add_bullet(doc, 'Hired in 2017 to start the gas trading business from scratch through developing and implementing growth '
               'strategies; successfully grew organically across all areas, including the client base (a global portfolio of +100 investors), '
               'risk appetite (from 0 VaR to $1m a week) and P&L ($7m growth)')
    add_bullet(doc, 'Part of a team of three Traders, grew the business from zero P&L to $7m in Gas and contributed to team-wide P&L of $50m')
    add_bullet(doc, 'Actively risk-managed the book with a probabilistic approach, structuring a portfolio of trades designed to have '
               'an attractive risk/reward profile under various scenarios. Successfully navigated extreme market conditions where daily '
               'volatility has been over 200%')
    add_bullet(doc, 'Generated alpha using linear and non-linear financial products across short- and long-term horizons through '
               'directional and relative value strategies; product, geographic and calendar spread trades and tail risk events')
    add_bullet(doc, 'Devised time-series forecasting models for European gas prices by integrating weather, utility, and macro indicators')
    add_bullet(doc, 'Engineered derivative trades leveraging Greeks, forward curves, and volatility surfaces to scale risk from 0 to $1M VaR weekly')
    add_bullet(doc, 'Streamlined risk reporting with real-time tools, cutting 1 hour/day of manual work and delivering insights to 100+ clients')

    add_sub_header(doc, 'Fundamental Analysis — Supply Related Modelling:')
    add_sub_bullet(doc, 'European Gas Production: For each gas production field, a range is derived by making separate assumptions '
                   'for future field development based on historical production and the future economics')
    add_sub_bullet(doc, 'European gas pipeline flows: Model the gas flow data by entry and exit points to track a real-time analysis '
                   'of supply levels versus historical levels and infrastructure utilisation rates')
    add_sub_bullet(doc, 'Maintenance: Forecasting flow disruptions through tracking maintenance schedules and pipeline activity')
    add_sub_bullet(doc, 'Gas storage: Utilising optimisation software which models the evolution of gas prices and derives the trading '
                   'strategy to optimise on the seasonality of gas prices; subject to the facility\'s constraints. A distribution of values '
                   'is generated from a Monte Carlo simulation of forward prices. The storage is modelled using advanced stochastics, '
                   'including Least Squares Monte Carlo techniques to capture the full optionality in gas storage facilities')
    add_sub_bullet(doc, 'LNG flows: Employing asset observation platforms which harness satellite imagery and machine learning applied '
                   'to images to track the LNG fleet and market dynamics and assess the performance of LNG plants and terminals')

    add_sub_header(doc, 'Fundamental Analysis — Demand Related Modelling:')
    add_sub_bullet(doc, 'Gas demand forecasting: Employing a linear regression model to determine the relationship between total demand '
                   'and a composite weather variable (an optimised combination of parameters of temperature and weather variables). A '
                   'scenario-based approach is used together with the model to forecast annual and daily demand')
    add_sub_bullet(doc, 'Simulate demand under different weather conditions by combining forecast daily demand models with historical '
                   'weather and applying statistical analysis to obtain 1 in 20 peak day demand forecasts')
    add_sub_bullet(doc, 'Weather: Incorporate data from sources including charts, maps, and satellite images and use algorithms to '
                   'quantify the potential impacts on gas markets of shifts in weather')
    add_sub_bullet(doc, 'Renewable Energy: Modelling and forecasting the capacity of renewable energy. Using weather models to predict '
                   'the wind and solar supply for power generation from renewable sources')
    add_sub_bullet(doc, 'Relative Fuel-Switching Costs: Incorporating Government carbon and climate policies into modelling each '
                   'country\'s fuel-switching cost curve to predict the power generation by fuel type')

    add_sub_header(doc, 'Quantitative Analysis:')
    add_sub_bullet(doc, 'Technical Analysis: Using price and price derivative data to generate technical signals to go overall long/short '
                   'in a trending market. Using methods of data visualisation, simulation, understanding signals and their robustness '
                   'and optimisation of strategies to enable rule-based trading')
    add_sub_bullet(doc, 'Developed trading strategies using quantitative models based on multiple data and model inputs, including mean '
                   'reversion; lead/lag relationships; roll differentials; rollover patterns; seasonality; contango and backwardation and '
                   'considering behaviour and positioning of other market participants')
    add_sub_bullet(doc, 'Option Spread Trading: Transformed NBP and JKM volatility pricing methodology, pricing it as a spread to TTF '
                   'vol, performing comparative analytics on their relative hub gas volatility surfaces and applying a relationship rule')
    add_sub_bullet(doc, 'Exotic Option Pricing: Worked closely with Quantitative Research and Exotic Oil Trading to calibrate Asian '
                   'options through regularly examining daily and monthly option pricing and understanding different volatility markets')

    add_sub_header(doc, 'Risk and Control Analysis:')
    add_sub_bullet(doc, 'Scenario analysis and portfolio stress testing: Measure the potential impact of discrete market events on '
                   'the portfolio using Value at Risk (VaR) methodologies')
    add_sub_bullet(doc, 'Backtesting strategies: Simulating how trading strategies evolve using historical data and employing '
                   'mathematics methods to draw statistical significance and validity')
    add_sub_bullet(doc, 'Risk charges: Collaborate with the Valuation Control Group to understand their methodology for calculating '
                   'risk price testing charges and explain the book\'s risk over weekly meetings')

    add_sub_header(doc, 'Additional Achievements:')
    add_bullet(doc, 'Created an interpolation function that embeds trends and seasonality to build the term structure of futures using illiquid spreads')
    add_bullet(doc, 'Transformed the linear and non-linear risk and pricing processes by developing a system to automate the daily practice '
               'of generating and assessing risk using VBA and cloud systems, eliminating the human error involved and reducing the computation '
               'time to seconds from twenty minutes — critical time on a Trading desk. Worked closely with Quant and Technology teams to update '
               'the bank risk and pricing systems allowing for more efficient pricing and risk management optics')
    add_bullet(doc, 'Collaborated with the Automated Trading Strategies Team to build the electronic trading platform to trade energy products')

    # ── JP Morgan — Graduate Analyst ──
    add_role_header(doc, 'JP Morgan', 'Graduate Analyst, Commodity Derivatives Marketing', 'London, UK', 'Oct 2014 – Dec 2016')
    add_body_text(
        doc,
        'Line of business: Youngest Analyst at JP Morgan working on the Commodities Marketing team.\n'
        'Reporting: Directly into the EMEA Head of Corporate Commodities Marketing.'
    )
    add_sub_header(doc, 'Main Duties:')
    add_bullet(doc, 'Understanding and providing the most suitable risk management solutions to corporate clients (consumers, producers, '
               'utilities, and refineries) using Energy, Metals and Agricultural Derivatives products in both a structured and flow environment')
    add_bullet(doc, 'Worked with Marketers and Traders to give indicative and live pricing to clients on derivative structures in the market needed')
    add_bullet(doc, 'Applying theoretical financial mathematics to assist in pricing exotics and complex structures as well as developing new '
               'products and structured deals, discussing pricing alternatives using a combination of derivatives to decide the optimal hedging strategy')
    add_bullet(doc, 'Researched hedging trends for EMEA, coordinated with APAC, NAMR and LATAM teams to produce a globally distributed report')
    add_bullet(doc, 'Worked closely with Senior Marketers and Traders to develop and structure bespoke derivatives products to suit corporate '
               'clients and to understand the risk management of each trade')
    add_bullet(doc, 'Performed analysis on existing and new products to explain their construction mechanics and performance attribution to clients')

    # ── JP Morgan — Summer Intern ──
    add_role_header(doc, 'JP Morgan', 'Summer and Off-Cycle Intern, Emerging Markets', 'London, UK', 'Jul 2014 – Sep 2014')
    add_bullet(doc, 'EM Sales, 8 weeks: shadowed Salespeople who marketed Asian, EMEA and LATAM Rates, FX, and Credit products to investors')
    add_bullet(doc, 'EM Trading, 2 weeks: shadowed Bond and Derivatives Traders and learnt about market-making, hedging and risk management')
    add_bullet(doc, 'Investor Client Management, 4 weeks: worked with 6 MDs to understand how they manage the bank\'s top 200 clients')

    # ── RBS — Summer Intern ──
    add_role_header(doc, 'Royal Bank of Scotland (RBS)', 'Summer Intern, CVA Trading and Prime Brokerage', 'London, UK', 'Jul 2013 – Sep 2013')
    add_bullet(doc, 'CVA Trading, 1 month: understood the computation of CVA for internal desks and risk management using CDS')
    add_bullet(doc, 'Prime Brokerage Sales, 1 month')

    # ── RBS — Spring Insight ──
    add_role_header(doc, 'Royal Bank of Scotland (RBS)', 'Spring Insight Intern in Markets', 'London, UK', 'Apr 2013')
    add_bullet(doc, 'Markets and Capital Market division overviews, including desk shadowing in FX Options Trading and FX Real Money Sales')

    # ── Earlier roles ──
    add_sub_header(doc, 'Earlier Roles:')

    add_role_header(doc, 'Next', 'Sales Associate', 'UK', 'Sep 2012 – Oct 2012')
    add_bullet(doc, 'Worked as part of the sales team to keep up to date with product knowledge and provide excellent customer service')

    add_role_header(doc, 'Ernst & Young', 'Intern, Assurance Department', 'Mauritius', 'Jul 2011 – Aug 2012')
    add_bullet(doc, 'Assisted in conducting the audit of a company in the manufacturing and industrial sector')

    add_role_header(doc, 'State Insurance Company of Mauritius', 'Intern, Marketing and Actuarial Departments', 'Mauritius', 'Jul 2011 – Aug 2011')
    add_bullet(doc, 'Marketing, 1 month: tracked media for the Company Intranet and researched marketing methods of similar companies')
    add_bullet(doc, 'Actuarial, 1 month: computed surrender/paid-up values and maturity benefits for life assurance policies')

    add_role_header(doc, 'Mauritius Broadcasting Corporation (MBC)', 'Screenwriter', 'Mauritius', 'Jan 2011 – Jun 2011')
    add_bullet(doc, 'Worked with the Producers and wrote scripts for the Knowledge Channel, a live children\'s television program')

    # ════════════════════════════════════════════════════════════════════════
    # 5. RESEARCH EXPERIENCE
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Research Experience')

    add_role_header(doc, 'Rice University & NASA', 'AI/ML Engineer', 'Houston, Texas', 'Jan 2025 – May 2025')
    p = add_bullet(doc, 'Achieved 93% precision in spacecraft localization via a deep learning pipeline using Python and PyTorch in Rice\'s D2K Lab')
    add_evidence_links(p, website_url='https://anekha.github.io/deep-learning/2025/03/28/NASA.html')
    add_bullet(doc, 'Designed and compared YOLO + Ellipse R-CNN vs. YOLO + geometry-based fitting on dual annotated/augmented datasets')
    add_bullet(doc, 'Fine-tuned each model separately using learning rate scheduling, early stopping, and experiment tracking with W&B')
    add_bullet(doc, 'Delivered CPU-optimized models (20s/image, <4GB RAM) and presented results at NASA Houston to support deployment')

    add_role_header(doc, 'Rice University & Stability AI', 'Computer Vision Researcher', 'Houston, Texas', 'Sep 2024 – May 2025')
    p = add_bullet(doc, 'ICLR 2026 paper: GIQ — Benchmarking 3D Geometric Reasoning of Vision Foundation Models with Simulated and Real Polyhedra')
    add_evidence_links(p, website_url='https://toomanymatts.github.io/giq-benchmark/')
    spacer = p.add_run(' ')
    spacer.font.size = Pt(10.5)
    spacer.font.name = 'Calibri'
    add_hyperlink(p, '[OpenReview]', 'https://openreview.net/forum?id=Uf8X57bQIr', color='8E44AD')
    add_bullet(doc, 'Collaborated with Stability AI\'s 3D Reconstruction Lab on dataset design, benchmarking, and evaluation')
    add_bullet(doc, 'Built and structured real/synthetic 3D datasets to test generalization of SOTA monocular and multi-view reconstruction models')
    add_bullet(doc, 'Analyzed model robustness and refined evaluation methods for LLM- and vision-based recognition under varied conditions')

    add_role_header(doc, 'ESKA International', 'Applied Research — Data Analytics in the Jewellery Industry', 'UK', '2021 – 2023')
    add_bullet(doc, 'Identified opportunities to apply data analytics to business decisions and improve processes in the jewellery industry, '
               'perceiving a gap between added commercial value and the concrete employment of innovative data science tools')
    add_bullet(doc, 'Conducted product and market research in the gemstone and jewellery industry to evaluate market opportunities')

    add_role_header(doc, 'Rice University', 'Ongoing Research — GaussianObject + MASt3R-SfM Integration', 'Houston, Texas', '2025 – Present')
    p = add_bullet(doc, 'Exploring how retrieval-based pose estimation (MASt3R-SfM) improves COLMAP-free Gaussian Splatting reconstruction pipelines')
    add_evidence_links(p, website_url='https://anekha.github.io/research/computer%20vision/3d/2025/02/28/MASt3RSfM.html')
    add_bullet(doc, 'Investigating sparse-view 3D reconstruction methods that combine Gaussian object representations with dense stereo matching')

    add_role_header(doc, 'JP Morgan', 'European Gas Research Analyst', 'London, UK', '2017 – 2021')
    add_bullet(doc, 'Alongside being the gas Trader, served as JP Morgan\'s European gas research analyst. Developed in-depth knowledge and '
               'wrote research pieces distributed to the broader team and clients and organised follow-up discussion calls')
    add_bullet(doc, 'Established strong relationships with external research providers and regularly held discussions with them to share '
               'opinions and challenge views')

    add_role_header(doc, 'University of Warwick', 'Operational Research Projects', 'UK', '2011 – 2014')
    add_sub_header(doc, 'Operational Research in Strategic Planning')
    add_bullet(doc, 'Explored the development and application of models and analytical techniques used to support a strategic planning '
               'process for the company Google Inc.')
    add_sub_header(doc, 'Decision Analysis')
    add_bullet(doc, 'Applied mathematical models and methods used in decision analysis and performance measurement to formulate '
               'business-related decision problems in a structured form suited to the solution using mathematical techniques. This project '
               'analysed the operations of a large grocery business with twenty-five small scale grocery stores. Interpreted the results and '
               'critically assessed limitations of the methods; used linear programming for modelling the decision-making processes')
    add_sub_header(doc, 'Simulation')
    add_bullet(doc, 'Used simulation software to virtually imitate the operations of a real-world process and predict future operational '
               'behaviours to give feedback and recommendations for improvement. Collected, modelled, and analysed data on the number of '
               'customers served at the busiest restaurant at the university and proposed an improved process using operational research '
               'methods for analysing complex operational industrial problems')

    # ════════════════════════════════════════════════════════════════════════
    # 6. PUBLICATIONS
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Publications')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('GIQ: Benchmarking 3D Geometric Reasoning of Vision Foundation Models with Simulated and Real Polyhedra')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_evidence_links(p, website_url='https://arxiv.org/abs/2506.08194')
    spacer = p.add_run(' ')
    spacer.font.size = Pt(10.5)
    spacer.font.name = 'Calibri'
    add_hyperlink(p, '[OpenReview]', 'https://openreview.net/forum?id=Uf8X57bQIr', color='8E44AD')
    spacer2 = p.add_run(' ')
    spacer2.font.size = Pt(10.5)
    spacer2.font.name = 'Calibri'
    add_hyperlink(p, '[Scholar]', 'https://scholar.google.com/citations?user=wpdCDjAAAAAJ&hl=en', color='2E86C1')
    add_body_text(doc, 'ICLR 2026')
    add_body_text(doc, 'Anekha Sokhal, Mateusz Michalkiewicz, Varun Jampani, Guha Balakrishnan, Tadeusz Michalkiewicz, '
                  'Piotr Pawlikowski, Mahsa Baktashmotlagh')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('International Consulting Program')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_evidence_links(p)
    add_body_text(doc, 'JP Morgan Investment Banking Blog (2020)')
    add_bullet(doc, 'Wrote an article published on the JP Morgan Investment Banking Blog on experience working with the Maharishi Institute')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('The Children\'s Knowledge Channel')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_evidence_links(p)
    add_body_text(doc, 'Mauritius Broadcasting Corporation (MBC), 2011')
    add_bullet(doc, 'Researched and wrote over ten scripts for the children\'s television programs aired on live television')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Young Writers, UK')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_evidence_links(p)
    add_bullet(doc, 'Five poems published across books, magazines and online for Young Writers UK')

    # ════════════════════════════════════════════════════════════════════════
    # 7. PROJECTS
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Projects')

    add_role_header(doc, 'COLMAP-Free Sparse 3D Reconstruction with Gaussian Splatting & MASt3R-SfM', '', '', 'Aug 2024 – Dec 2024')
    add_bullet(doc, 'Built a sparse-view 3D pipeline by integrating Gaussian Splatting and MASt3R for pose estimation and surface refinement')
    add_bullet(doc, 'Improved reconstruction fidelity via dense pose fusion and iterative SfM refinement, reducing reprojection error by over 60%')
    add_bullet(doc, 'Achieved high-quality 3D assets from just <20 input images, enabling real-time rendering without photogrammetry tools')

    add_role_header(doc, 'Face-Based Personalized Jewelry Recommendation System', '', '', 'Jan 2024 – May 2024')
    add_bullet(doc, 'Achieved 90%+ accuracy on 5,000-image Kaggle dataset; enhanced user experience through visually tailored outputs')
    add_bullet(doc, 'Launched a Streamlit app using MTCNN, OpenCV, and a CNN to personalise jewellery suggestions by face shape and skin tone')

    add_role_header(doc, 'Student Performance Forecaster Model', '', '', 'Mar 2023')
    add_bullet(doc, 'Built a Gradient Boosting Classifier to predict academic outcomes using socio-economic data; achieved 96% accuracy')
    add_bullet(doc, 'Deployed a Streamlit dashboard integrating feature engineering, EDA, and model interpretability for educator insights')

    # ════════════════════════════════════════════════════════════════════════
    # 8. AWARDS & RECOGNITION
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Awards & Recognition')
    p = add_bullet(doc, 'UK-US Fulbright Elsevier Data Analytics Award (2023)')
    add_evidence_links(p, website_url='https://fulbright.org.uk/people-search/anekha-sokhal/')
    p = add_bullet(doc, 'IBM Masters Fellowship Awards in AI — Nominee (2024)')
    add_evidence_links(p)
    p = add_bullet(doc, 'Women in AI North America Awards — Finalist (2025)')
    add_evidence_links(p, website_url='https://www.womeninai.co/post/finalists-announced-for-women-in-ai-awards-north-america-2025')
    p = add_bullet(doc, 'Lilie Accelerator — Selected for Rice\'s Liu Idea Lab Summer Venture Studio as 1 of 8 startups (2024)')
    add_evidence_links(p, website_url='https://news.rice.edu/news/2024/rice-entrepreneurship-lab-unveils-teams-2024-lilie-summer-venture-studio-accelerator')
    p = add_bullet(doc, 'Ranked 11th in Mauritius in Cambridge International Examinations A-Levels Girls\' Science Side (2010)')
    add_evidence_links(p)
    p = add_bullet(doc, 'Ranked 1st in Mauritius in Economics in Cambridge International Examinations O-Levels Girls\' Side (2008)')
    add_evidence_links(p)
    p = add_bullet(doc, 'Best Delegate and Delegation for representing Georgia in the Model United Nations Conference (2010)')
    add_evidence_links(p)
    p = add_bullet(doc, 'Work Experience Bursary, University of Warwick (2012)')
    add_evidence_links(p)
    p = add_bullet(doc, '3rd prize, Code First: Girls web development competition (2017)')
    add_evidence_links(p)

    # ════════════════════════════════════════════════════════════════════════
    # 9. SPEAKER ENGAGEMENTS & JUDGING
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Speaker Engagements & Judging')
    p = add_bullet(doc, 'NRLC 2026 Judge — Judge at the Napier Rice Launch Challenge (Mar 2026)')
    add_evidence_links(p, website_url='https://entrepreneurship.rice.edu/napier-rice-launch-challenge')
    add_bullet(doc, 'Mercury Fund Day at the Ion — Panel speaker on "The New FDE Model: How to Deploy with Scale in Legacy Industries" (Feb 2026)')
    p = add_bullet(doc, 'Lilie Founders Journey — Guest speaker in the Founders Journey class at Rice Entrepreneurship, filmed as a podcast episode (Feb 2026)')
    add_evidence_links(p, website_url='https://entrepreneurship.rice.edu/')
    add_bullet(doc, 'Rice MDS Student–Alumni Panel — Panelist at the Master of Data Science student–alumni discussion, invited by Dr. Marmar Orooji (Dec 2025)')
    p = add_bullet(doc, 'Houston AI Club — Invited speaker at Houston\'s largest AI meetup community (May 2025)')
    add_evidence_links(p, website_url='https://www.houstonaiclub.com/')
    add_bullet(doc, 'Rice MBA Summer Startup Sync — Guest speaker presenting to Rice MBA students on startups and venture building (May 2025)')
    p = add_bullet(doc, 'Generative AI & ML in the Enterprise — Presented applications of generative AI and computer vision in e-commerce at the Data Science Salon, Austin (Feb 2025)')
    add_evidence_links(p, website_url='https://www.datascience.salon/austin/')
    add_bullet(doc, 'Fulbright "Connect with Your Community" Webinar — Panelist in a Fulbright Program webinar on community engagement. Recognized on the Fulbright Program Star list (Oct 2024)')
    add_bullet(doc, 'Houston Fulbright Enrichment Seminar — Selected to attend enrichment seminar in Houston (Feb 2024)')
    add_bullet(doc, 'JewelVision Demo Day — Pitched to 100+ investors and researchers (2024)')

    # ════════════════════════════════════════════════════════════════════════
    # 10. MEDIA & PODCASTS
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Media & Podcasts')
    p = add_bullet(doc, 'AI After Dark Podcast — Guest on the AI After Dark podcast, hosted by Alex Gras at Collide and Mercury Fund (Feb 2026)')
    add_evidence_links(p, website_url='https://app.collide.io/content/')
    add_bullet(doc, 'Lilie Founders Journey Podcast — Featured episode from the Founders Journey class at Rice Entrepreneurship (Feb 2026)')
    p = add_bullet(doc, 'Don\'t Panic Podcast — Guest episode (Feb 2023)')
    add_evidence_links(p, website_url='https://dontpanicpodcast.buzzsprout.com/2029474/episodes/12279537')

    # ════════════════════════════════════════════════════════════════════════
    # 11. ADVISORY
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Advisory')
    add_role_header(doc, 'Quicksilver AI Labs, Inc.', 'AI Advisor', '', 'Dec 2025 – Present')
    p = add_bullet(doc, 'Providing strategic AI and technical guidance to an early-stage AI venture fund')
    add_evidence_links(p)

    # ════════════════════════════════════════════════════════════════════════
    # 12. ADDITIONAL CONSULTING & SOCIAL IMPACT
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Additional Consulting & Social Impact')

    add_role_header(doc, 'Unrest', 'Business Consultant', '', 'Mar 2022 – Apr 2022')
    add_bullet(doc, 'Unrest is a mission-driven start-up accelerator program that develops budding entrepreneurs who want to launch '
               'mission-driven businesses and supports them to ensure they make a global impact. Works with thirty future Impact Unicorns '
               'each year and takes them from MVP to an operational business that\'s ready to fundraise within four months')
    add_bullet(doc, 'Conducted a comprehensive analysis of the current strategic plan, financial forecasts, market positioning and competitive landscape')
    add_bullet(doc, 'Worked closely with the Founder to establish future-facing insights, implications, and recommendations, including '
               'identifying potential opportunities and risks involved in each strategy')

    add_role_header(doc, 'JP Morgan Virtual Service Corps / Maharishi Institute', 'Business Consultant, International Consulting Program', 'South Africa (Remote)', 'Sep 2019 – Mar 2020')
    add_bullet(doc, 'Selected as part of a team of six consultants globally working together with Pyxera Global and an educational '
               'non-profit in South Africa, the Maharishi Invincibility Institute, the world\'s first self-funded university, as part of '
               'JP Morgan\'s International Consulting Program with the aim of developing a financial expansion plan to meet development '
               'goals for the social enterprise')
    add_bullet(doc, 'Created an integrated dashboard to track and forecast the NGO\'s OPEX and CAPEX, as well as identify critical '
               'success factors and areas for optimisation and innovations to its model')

    add_role_header(doc, 'ThoughtWorks & HumanityX', 'Project Consultant, "Building an Equitable Tech Future"', 'London, UK', 'Oct 2017 – Dec 2017')
    add_bullet(doc, 'Working with a team of technology consultants and engineers to research the social impact of tech trends and how '
               'they can disrupt the charity sector by streamlining donations using technology to manage costs and processes')

    add_role_header(doc, 'Personal Project', 'Project Manager — House Renovation', 'London, UK', 'Jul 2018 – Feb 2019')
    add_bullet(doc, 'Coordinating and scheduling the activities of tradespeople to renovate a 3-bedroom, 2-bathroom house from scratch, '
               'including all electrical, plumbing, fire safety and cosmetic work; completed the renovation project on schedule and within budget')

    add_role_header(doc, 'Warwick International Ventures (WIV)', 'Participant — Mumbai Trip', '', 'Jul 2013')
    add_bullet(doc, 'Selected to travel to Mumbai to meet with executives from Deutsche Bank, Bloomberg, and Citi Group India, to develop '
               'a more comprehensive worldview, and understand financial practices within an emerging market')

    add_role_header(doc, 'Model United Nations Conference', 'Delegate representing Georgia in the Economic Commission', '', 'Jul 2010')
    add_bullet(doc, 'Increased international understanding and developed negotiation skills')

    # ════════════════════════════════════════════════════════════════════════
    # 13. TEACHING & MENTORING
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Teaching & Mentoring')

    add_role_header(doc, 'JP Morgan', 'Commodities Analyst Training', 'London, UK', '2015 – 2021')
    add_bullet(doc, 'Responsible for the formal and informal training of new graduates on the Commodities desk. The training consisted '
               'of daily work shadowing and detailed explanations of the work, including the technicalities of the market, pricing, client '
               'coverage and evaluation of work. Personally trained five individual graduates and over ten interns during time at JP Morgan')

    add_role_header(doc, 'JP Morgan', 'Workshops to Summer Interns', 'London, UK', '2015 – 2021')
    add_bullet(doc, 'Active member of the recruitment team and gave regular workshops on Commodities and Energy Trading to the internship class')

    add_role_header(doc, 'JP Morgan', 'Mentor to Summer Interns', 'London, UK', 'Jun 2015 – Sep 2017')
    add_bullet(doc, 'Weekly meetings to work on goal setting, providing guidance and support to mentees')

    add_role_header(doc, 'University of Warwick', 'Mentor to First Year MORSE Students', 'UK', 'Sep 2013 – Jun 2014')
    add_bullet(doc, 'Shared experience and advice with new students with the aim of positively impacting their future')

    add_role_header(doc, 'University of Warwick', 'Orientation Helper, International Office', 'UK', 'Sep 2013')
    add_bullet(doc, 'Demonstrated enthusiasm, dedication, and friendliness during the welcoming week for over a thousand new international students')

    add_role_header(doc, 'Rice University', 'Fulbright Leadership Committee', 'Houston, Texas', 'Sep 2024 – May 2025')
    add_bullet(doc, 'Planned calendar of activities, coordinated events, and advised the GPS office on programming for 100+ international scholars')

    add_role_header(doc, 'Doerr Institute, Rice University', 'Women in Leadership Program', 'Houston, Texas', '2024')
    add_bullet(doc, 'Completed executive coaching and leadership development')

    # ════════════════════════════════════════════════════════════════════════
    # 14. VOLUNTEERING
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Volunteering')

    add_role_header(doc, 'Crisis at Christmas', 'Volunteer — UK homelessness charity', 'London, UK', 'Dec 2017')
    add_bullet(doc, 'Part of a team of volunteers who engaged with and supported guests whilst ensuring the centres ran smoothly over Christmas')

    add_role_header(doc, '"Beyond the Bottom Line" Social Enterprise Conference, University of Warwick',
                    'Speakers Team Coordinator', 'UK', 'Sep 2012 – Jan 2013')
    add_bullet(doc, 'Part of the Organising Committee for the first student-led social enterprise conference at Warwick University')
    add_bullet(doc, 'Researched, communicated, and managed nine high-profile speakers')

    add_role_header(doc, 'Warwick Jailbreak for the charity "Practical Action"', 'Volunteer & Team Leader', 'UK / Germany', 'Nov 2012')
    add_bullet(doc, 'The team leader of three; hitchhiked close to five-hundred miles to Cologne, Germany, in 36 hours without any money')

    add_role_header(doc, 'Queen Elizabeth College', 'Founder & President of the Environment Club', 'Mauritius', 'Jan 2009 – Nov 2010')
    add_bullet(doc, 'Founded the Environment Club and organised the weekly recycling of plastic PET bottles on campus')

    # ════════════════════════════════════════════════════════════════════════
    # 15. PROFESSIONAL DEVELOPMENT
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Professional Development')

    dev_items = [
        ('The Data Science Course 2022: Complete Data Science Bootcamp', '365 Careers / Udemy', 'May 2022',
         'A complete Data Science course covering mathematics, statistics, Python, statistics in Python, machine & deep learning'),
        ('ESG Investing', 'FitchLearning', 'May 2022',
         'Course on ESG market and engagement, ESG investment analysis and integrating ESG into portfolio management'),
        ('Mathematics for Machine Learning Specialization', 'Imperial College London', 'May 2022',
         'Mathematics course on Linear Algebra, Multivariate Calculus and Dimensionality Reduction with Principal Component Analysis'),
        ('Introduction to Data & SQL', 'Code First: Girls', 'May 2022', None),
        ('Introduction to Python & Apps', 'Code First: Girls', 'Mar 2022 – May 2022', None),
        ('Data Science Beginner\'s Tutorial', 'freeCodeCamp', 'Apr 2022', None),
        ('Social Enterprise Workshop for emerging entrepreneurs', 'Hatch Launchpad Programme', 'Nov 2017',
         'Aimed at helping entrepreneurs who want to start and grow a sustainable business through workshops, coaching and mentoring'),
        ('JP Morgan Associate Programme in New York', 'JP Morgan', 'Apr 2017',
         'A global training program delivered to newly-promoted Associates aimed at developing their network and professional skills'),
        ('Advanced Python Web Development Course', 'Code First: Girls', 'Jan 2017 – Mar 2017', None),
        ('Beginners HTML/CSS Web Development Course', 'Code First: Girls', 'Sep 2016 – Jan 2017',
         'Introduction to the use of HTML & CSS. Awarded 3rd prize in course web development competition'),
        ('Introduction to Corporate Finance', 'JP Morgan', 'Jun 2017',
         'Covering financial analysis of organisations and the industry performance considering various financial parameters'),
        ('JP Morgan Analyst Training Programme in New York', 'JP Morgan', 'Aug 2015 – Oct 2015',
         'Intensive training programme across all financial asset classes with weekly examinations (87% average achieved)'),
        ('Women in Networking Leadership Programme', '', 'May 2011 – Aug 2011',
         'Youngest graduate on a course to empower women to become leaders and asked to deliver a closing ceremony speech'),
    ]

    for name, provider, dates, description in dev_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run_name = p.add_run(name)
        run_name.bold = True
        run_name.font.size = Pt(10.5)
        run_name.font.name = 'Calibri'
        if provider:
            run_prov = p.add_run(f', {provider}')
            run_prov.font.size = Pt(10.5)
            run_prov.font.name = 'Calibri'
        run_date = p.add_run(f'\t{dates}')
        run_date.font.size = Pt(10.5)
        run_date.font.name = 'Calibri'
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        if description:
            add_bullet(doc, description)

    # ════════════════════════════════════════════════════════════════════════
    # 16. SKILLS & LANGUAGES
    # ════════════════════════════════════════════════════════════════════════
    add_section_heading(doc, 'Skills & Languages')

    skills = [
        ('Programming', 'Python, SQL, R, VBA, HTML, CSS, Spark, Hadoop, Bash'),
        ('ML/AI Frameworks', 'PyTorch, scikit-learn, OpenCV, Hugging Face'),
        ('3D Vision', 'NeRF, COLMAP, Gaussian Splatting'),
        ('Data Science', 'pandas, NumPy, SciPy'),
        ('Cloud & Platforms', 'AWS, Modal, Google Colab'),
        ('Domains', 'Deep Learning, Large Language Models (LLMs), Computer Vision, 3D Reconstruction'),
        ('Finance Tools', 'Bloomberg, Reuters'),
        ('Business Tools', 'Sage Accounting, Quickbooks'),
        ('Languages', 'English (Native), French (Fluent), Mauritian Creole (Fluent), Hindi (Intermediate)'),
    ]

    for category, items in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_cat = p.add_run(f'{category}: ')
        run_cat.bold = True
        run_cat.font.size = Pt(10.5)
        run_cat.font.name = 'Calibri'
        run_items = p.add_run(items)
        run_items.font.size = Pt(10.5)
        run_items.font.name = 'Calibri'

    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ANEKHA_SOKHAL_O1_CV.docx')
    doc.save(output_path)
    print(f'CV saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_cv()
